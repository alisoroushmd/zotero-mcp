"""MCP server exposing Zotero tools via FastMCP."""

from __future__ import annotations

import atexit
import functools
import json
import logging
import os
import re
import tempfile
import threading
import time

import httpx
from fastmcp import FastMCP

from typing import TYPE_CHECKING

from zotero_mcp.capabilities import check_capabilities, format_status
from zotero_mcp.config import get_config
from zotero_mcp.local_client import LocalClient
from zotero_mcp.web_client import WebClient, _is_preprint_doi

if TYPE_CHECKING:
    from zotero_mcp.knowledge_graph import KnowledgeGraph

logger = logging.getLogger(__name__)

mcp = FastMCP(
    "zotero",
    instructions=(
        "Zotero MCP server. All tools work with just API credentials "
        "(ZOTERO_API_KEY + ZOTERO_USER_ID). If Zotero desktop is also "
        "running, reads are faster via the local API. "
        "Call server_status to check available modes. "
        "Use the prompts (literature_audit, build_and_explore, add_and_verify, "
        "extract_entities) for guided multi-step workflows."
    ),
)

_local: LocalClient | None = None
_local_failed_at: float | None = None  # time.monotonic() when probe last failed
_LOCAL_RETRY_INTERVAL = 300.0  # retry local API probe every 5 minutes
_web: WebClient | None = None
_init_lock = threading.Lock()

# Temp file tracking for cleanup at exit.
_temp_files: list[str] = []
_temp_lock = threading.Lock()


def _register_temp_file(path: str) -> None:
    """Register a temporary file for cleanup at process exit."""
    with _temp_lock:
        _temp_files.append(path)


def _cleanup_temp_files() -> None:
    """Remove all registered temporary files."""
    with _temp_lock:
        for path in _temp_files:
            try:
                os.unlink(path)
            except OSError:
                pass
        _temp_files.clear()


atexit.register(_cleanup_temp_files)

_ZOTERO_KEY_RE = re.compile(r"^[A-Za-z0-9]+$")

# Directories allowed for file read/write operations.
_ALLOWED_PATH_ROOTS: list[str] = []


def _get_allowed_path_roots() -> list[str]:
    """Return the set of directories allowed for file I/O.

    Includes the current working directory, user home, and system temp dir.
    Computed lazily and cached.
    """
    if not _ALLOWED_PATH_ROOTS:
        import pathlib

        candidates = [
            pathlib.Path.cwd(),
            pathlib.Path.home(),
            pathlib.Path(tempfile.gettempdir()),
        ]
        for c in candidates:
            try:
                _ALLOWED_PATH_ROOTS.append(str(c.resolve()))
            except OSError:
                pass
    return _ALLOWED_PATH_ROOTS


def _validate_path(file_path: str, name: str = "path") -> str:
    """Validate that a file path is within allowed directories.

    Resolves the path (following symlinks) and checks it falls under
    the current working directory, user home, or system temp directory.

    Returns the resolved absolute path string.
    """
    import pathlib

    resolved = str(pathlib.Path(file_path).resolve())
    allowed = _get_allowed_path_roots()
    if not any(
        resolved.startswith(root + os.sep) or resolved == root for root in allowed
    ):
        raise ValueError(
            f"{name} must be within the working directory, home directory, "
            f"or temp directory. Got: {file_path!r}"
        )
    return resolved


def _validate_key(value: str, name: str = "key") -> None:
    """Validate a Zotero item/collection key."""
    if not value or not value.strip():
        raise ValueError(f"{name} must not be empty")
    if not _ZOTERO_KEY_RE.match(value.strip()):
        raise ValueError(f"{name} must be alphanumeric, got: {value!r}")


def _clamp_limit(value: str | int, lo: int = 1, hi: int = 100) -> int:
    """Clamp a limit parameter to a safe range."""
    return max(lo, min(hi, int(value)))


def _get_local() -> LocalClient:
    """Lazy-initialize the local client with TTL-based failure caching.

    Retries the probe every _LOCAL_RETRY_INTERVAL seconds so that starting
    Zotero desktop mid-session is automatically picked up.
    """
    global _local, _local_failed_at
    now = time.monotonic()
    if (
        _local_failed_at is not None
        and (now - _local_failed_at) < _LOCAL_RETRY_INTERVAL
    ):
        raise RuntimeError("Local API unavailable (cached)")
    if _local is None:
        with _init_lock:
            now = time.monotonic()
            if (
                _local_failed_at is not None
                and (now - _local_failed_at) < _LOCAL_RETRY_INTERVAL
            ):
                raise RuntimeError("Local API unavailable (cached)")
            if _local is None:
                try:
                    _local = LocalClient()
                    _local_failed_at = None  # Reset on successful probe
                    logger.info("Local Zotero API connected")
                except RuntimeError:
                    _local_failed_at = time.monotonic()
                    logger.info(
                        "Local Zotero API unavailable — will retry in %.0fs",
                        _LOCAL_RETRY_INTERVAL,
                    )
                    raise
    return _local


def _get_web() -> WebClient:
    """Lazy-initialize the web client (thread-safe)."""
    global _web
    if _web is not None:
        return _web

    with _init_lock:
        if _web is not None:
            return _web
        cfg = get_config()
        if not cfg.has_web_api:
            raise RuntimeError(
                f"Cloud CRUD mode requires {', '.join(cfg.missing_web_vars)}. "
                f"Get your API key at https://www.zotero.org/settings/keys"
            )
        api_key = cfg.zotero_api_key
        user_id = cfg.zotero_user_id
        # Try to attach local client for faster reads, but don't fail without it
        local = None
        try:
            local = _get_local()
        except RuntimeError:
            pass  # Zotero desktop not running — web client will use web reads
        _web = WebClient(api_key=api_key, user_id=user_id, local_client=local)
        return _web


def _parse_list_param(value: str | list | None) -> list | None:
    """Parse a parameter that may be a JSON string, list, or None."""
    if value is None:
        return None
    if isinstance(value, list):
        return value
    try:
        parsed = json.loads(value)
        return parsed if isinstance(parsed, list) else [value]
    except (json.JSONDecodeError, TypeError):
        return [value]  # Treat bare string as single-item list


def _error_response(code: str, message: str, **extra) -> dict:
    """Build a structured error dict for MCP tool return values.

    Args:
        code: Short machine-readable error code (e.g. "invalid_key").
        message: Human-readable explanation.
        **extra: Additional fields to include in the response.

    Returns:
        Dict with "error", "message", and any extra keys.
    """
    return {"error": code, "message": message, **extra}


def _handle_tool_errors(fn):
    """Decorator that converts common exceptions to structured JSON error responses.

    Catches ValueError (bad input), RuntimeError (unavailable/config), and
    httpx.HTTPStatusError (API errors), and returns a JSON-encoded error dict
    instead of raising so the LLM receives a readable error message.
    """

    @functools.wraps(fn)
    def _wrapper(*args, **kwargs):
        try:
            return fn(*args, **kwargs)
        except ValueError as exc:
            return json.dumps(_error_response("invalid_input", str(exc)))
        except httpx.HTTPStatusError as exc:
            return json.dumps(
                _error_response(
                    "api_error",
                    f"Zotero API returned {exc.response.status_code}",
                    status_code=exc.response.status_code,
                )
            )
        except RuntimeError as exc:
            return json.dumps(_error_response("unavailable", str(exc)))

    return _wrapper


# -- Status tool --


@mcp.tool(
    description=(
        "Check which Zotero MCP operating modes are available and get fix instructions "
        "for any that are misconfigured. Use this first when tools return 'unavailable' errors "
        "or when starting a new session to verify connectivity."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def server_status() -> str:
    """Probe Zotero services and report available modes and tools."""
    caps = check_capabilities()
    return json.dumps(format_status(caps), ensure_ascii=False)


# -- Read tools (web API primary, local API fallback) --


def _read_local_or_web(local_method: str, *args, **kwargs):
    """Try local API first (faster), fall back to web API for reads."""
    try:
        local = _get_local()
        return getattr(local, local_method)(*args, **kwargs)
    except RuntimeError:
        try:
            return getattr(_get_web(), local_method)(*args, **kwargs)
        except httpx.TimeoutException as exc:
            raise RuntimeError(
                f"Zotero Web API timed out ({exc.__class__.__name__}). "
                "Try a more specific query, reduce the limit, or start "
                "Zotero desktop for faster local searches."
            ) from exc


@mcp.tool(
    description=(
        "Search Zotero library items by keyword. Use this when the user asks to "
        "find papers, look up references, or search their library. Supports title, "
        "author, and tag matching. Filter by item_type (e.g. 'journalArticle', "
        "'book') or tag (exact tag name)."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def search_items(
    query: str,
    limit: str | int = 25,
    item_type: str = "",
    tag: str = "",
) -> str:
    """Search for items by keyword. Excludes attachments and notes."""
    results = _read_local_or_web(
        "search_items",
        query,
        _clamp_limit(limit),
        item_type=item_type or None,
        tag=tag or None,
    )
    return json.dumps(results, ensure_ascii=False)


@mcp.tool(
    description=(
        "Get detailed metadata for a single Zotero item by its key. Use this when "
        "you need full bibliographic details (title, authors, DOI, abstract, dates) "
        "for a specific item. Set format='bibtex' for BibTeX export."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def get_item(item_key: str, format: str = "json") -> str:
    """Get full metadata or BibTeX for one item by its key."""
    _validate_key(item_key, "item_key")
    result = _read_local_or_web("get_item", item_key.strip(), fmt=format)
    if isinstance(result, str):
        return result
    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "List all collections (folders) in the Zotero library with their keys, names, "
        "parent relationships, and item counts. Use this when the user asks about their "
        "library organization or wants to browse/find a collection."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def get_collections() -> str:
    """Returns flat list of collections with key, name, parent, and item count."""
    results = _read_local_or_web("get_collections")
    return json.dumps(results, ensure_ascii=False)


@mcp.tool(
    description=(
        "Get all notes attached to a Zotero item. Use this when the user wants to "
        "read their annotations, reading notes, or comments on a paper."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def get_notes(parent_key: str) -> str:
    """Get all notes attached to a parent item."""
    _validate_key(parent_key, "parent_key")
    results = _read_local_or_web("get_notes", parent_key.strip())
    return json.dumps(results, ensure_ascii=False)


@mcp.tool(
    description=(
        "List file attachments (PDFs, etc.) on a Zotero item with availability status. "
        "Use this to check whether a paper has a PDF before trying to read it. "
        "Returns availability: stored_remote_available, stored_local_available, "
        "linked_local_available, or metadata_only."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def get_item_attachments(parent_key: str) -> str:
    """Get attachments for a parent item with availability classification."""
    _validate_key(parent_key, "parent_key")
    attachments = _read_local_or_web(
        "get_children", parent_key.strip(), item_type="attachment"
    )

    link_mode_map = {
        "imported_url": "stored_remote_available",
        "imported_file": "stored_local_available",
        "linked_file": "linked_local_available",
        "linked_url": "linked_local_available",
    }

    results = []
    for att in attachments:
        link_mode = att.get("linkMode", "")
        results.append(
            {
                "key": att.get("key", ""),
                "filename": att.get("filename", ""),
                "contentType": att.get("contentType", ""),
                "availability": link_mode_map.get(link_mode, "metadata_only"),
            }
        )
    return json.dumps(results, ensure_ascii=False)


@mcp.tool(
    description=(
        "Get the full text of a paper in the library. Routes to the best available source: "
        "PubMed Central, local PDF, web PDF download, or free open-access PDF. "
        "Use this when the user wants to read a paper's content. "
        "Set extract_text=true to extract and return the text inline; "
        "otherwise returns a file path or PMCID for the caller to read."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def get_pdf_content(item_key: str, extract_text: bool = False) -> str:
    """Route to the best available content source for a Zotero item.

    Args:
        item_key: Zotero item key.
        extract_text: If True, extract text from the PDF and return it inline.

    Returns:
        JSON with content_source and the relevant identifier or path.
    """
    _validate_key(item_key, "item_key")
    item_key = item_key.strip()

    # Read item metadata
    item = _read_local_or_web("get_item", item_key)
    if isinstance(item, str):
        return json.dumps(
            {
                "item_key": item_key,
                "content_source": "error",
                "message": "Could not read item metadata.",
            }
        )

    doi = item.get("DOI", "")
    extra = item.get("extra", "")
    url = item.get("url", "")

    def _maybe_extract(source, result_dict):
        """If extract_text is requested, extract text from PDF source and add to result."""
        if not extract_text:
            return result_dict
        from zotero_mcp.text_extractor import extract_text_from_pdf

        text = extract_text_from_pdf(source)
        if text:
            result_dict["content_source"] = "extracted_text"
            result_dict["text"] = text
            result_dict["page_count"] = text.count("\n\n") + 1
            result_dict["char_count"] = len(text)
            result_dict.pop("message", None)
        return result_dict

    # Step 1: Check for PMCID via PMID (skip when extracting — need actual PDF)
    if not extract_text:
        pmid_match = re.search(r"PMID:\s*(\d+)", extra)
        if pmid_match:
            pmid = pmid_match.group(1)
            try:
                pmcid = _get_web().resolve_pmid_to_pmcid(pmid)
                if pmcid:
                    return json.dumps(
                        {
                            "item_key": item_key,
                            "content_source": "pmc",
                            "pmcid": pmcid,
                            "pmid": pmid,
                            "message": "Use PubMed MCP get_full_text_article(pmcid)",
                        }
                    )
            except Exception as exc:
                logger.warning(
                    "PMCID lookup failed for item %s PMID %s: %s", item_key, pmid, exc
                )

    # Step 2: Check for PDF attachments
    try:
        children = _read_local_or_web("get_children", item_key, item_type="attachment")
    except Exception as exc:
        logger.warning("Failed to list attachments for %s: %s", item_key, exc)
        children = []

    pdf_attachments = [c for c in children if c.get("contentType") == "application/pdf"]

    if pdf_attachments:
        att = pdf_attachments[0]
        att_key = att.get("key", "")

        # Step 3: Try local file path (fastest)
        try:
            local = _get_local()
            local_path = local.get_attachment_path(att_key)
            if local_path:
                result = {
                    "item_key": item_key,
                    "content_source": "local_pdf",
                    "pdf_path": local_path,
                    "attachment_key": att_key,
                    "message": "Read this PDF path",
                }
                return json.dumps(_maybe_extract(local_path, result))
        except Exception as exc:
            logger.warning(
                "Local attachment path lookup failed for %s: %s", att_key, exc
            )

        # Step 4: Download from web API
        try:
            web = _get_web()
            pdf_bytes = web.download_attachment(att_key)
            if extract_text:
                result = {
                    "item_key": item_key,
                    "content_source": "web_pdf",
                    "attachment_key": att_key,
                }
                return json.dumps(_maybe_extract(pdf_bytes, result))
            tmp = tempfile.NamedTemporaryFile(
                prefix="zotero_mcp_", suffix=".pdf", delete=False
            )
            try:
                tmp.write(pdf_bytes)
                tmp.close()
            except Exception:
                tmp.close()
                os.unlink(tmp.name)
                raise
            _register_temp_file(tmp.name)
            return json.dumps(
                {
                    "item_key": item_key,
                    "content_source": "web_pdf",
                    "pdf_path": tmp.name,
                    "attachment_key": att_key,
                    "message": "Read this PDF path",
                }
            )
        except Exception as exc:
            logger.warning(
                "Web PDF download failed for attachment %s: %s", att_key, exc
            )

    # Step 5: No stored PDF — try free PDF via DOI (Unpaywall / PMC / bioRxiv)
    if doi:
        try:
            pdf_bytes, _, source = _get_web()._download_free_pdf(doi)
            if pdf_bytes:
                if extract_text:
                    result = {
                        "item_key": item_key,
                        "content_source": f"free_pdf_{source}",
                        "doi": doi,
                    }
                    return json.dumps(_maybe_extract(pdf_bytes, result))
                tmp = tempfile.NamedTemporaryFile(
                    prefix="zotero_mcp_", suffix=".pdf", delete=False
                )
                try:
                    tmp.write(pdf_bytes)
                    tmp.close()
                except Exception:
                    tmp.close()
                    os.unlink(tmp.name)
                    raise
                _register_temp_file(tmp.name)
                return json.dumps(
                    {
                        "item_key": item_key,
                        "content_source": f"free_pdf_{source}",
                        "pdf_path": tmp.name,
                        "doi": doi,
                        "message": "Read this PDF path",
                    }
                )
        except Exception as exc:
            logger.warning("Free PDF download failed for DOI %s: %s", doi, exc)

    # Step 6: No PDF found anywhere, return DOI/URL fallback
    result: dict = {
        "item_key": item_key,
        "content_source": "not_found",
        "message": "No PDF attached or available open-access. Try DOI or ask user for the file.",
    }
    if doi:
        result["doi"] = doi
    if url:
        result["url"] = url
    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Check whether papers have been retracted or corrected. Uses CrossRef "
        "(authoritative) and OpenAlex. Use this when the user asks about paper "
        "validity, before citing papers, or as part of a literature audit. "
        "Accepts one or more item keys."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def check_retractions(item_keys: str | list[str]) -> str:
    """Batch check items for retractions and corrections.

    Args:
        item_keys: Single key or list of Zotero item keys.

    Returns:
        JSON with per-item retraction/correction status and summary counts.
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed

    from zotero_mcp.openalex_client import OpenAlexClient

    keys = _parse_list_param(item_keys) or []
    if not keys:
        raise ValueError("item_keys must not be empty")
    for k in keys:
        _validate_key(k, "item_key")

    web = _get_web()
    openalex = OpenAlexClient()

    results = []
    retracted_count = 0
    corrected_count = 0

    def _check_one(key: str) -> dict:
        item = web.get_item(key.strip())
        if isinstance(item, str):
            return {"key": key, "error": "Could not read item"}

        doi = item.get("DOI", "")
        title = item.get("title", "")

        entry: dict = {"key": key, "title": title, "retracted": False}

        if not doi:
            entry["warning"] = "No DOI — cannot check retraction status"
            return entry

        entry["doi"] = doi

        # CrossRef (authoritative for retractions)
        crossref = web.check_crossref_updates(doi)
        if crossref["has_retraction"]:
            entry["retracted"] = True
            entry["retraction_doi"] = crossref["retraction_doi"]
            entry["retraction_date"] = crossref["retraction_date"]
        if crossref["corrections"]:
            entry["corrections"] = crossref["corrections"]

        # OpenAlex (broader context + citation count)
        oa_work = openalex.get_work(doi)
        if oa_work:
            cited_by = oa_work.get("cited_by_count", 0)
            if cited_by:
                entry["cited_by_count"] = cited_by
            # OpenAlex retraction flag as backup
            if oa_work.get("is_retracted") and not entry["retracted"]:
                entry["retracted"] = True
                entry["retraction_source"] = "openalex"

        return entry

    with ThreadPoolExecutor(max_workers=3) as pool:
        futures = {pool.submit(_check_one, k): k for k in keys}
        for future in as_completed(futures):
            entry = future.result()
            results.append(entry)
            if entry.get("retracted"):
                retracted_count += 1
            if entry.get("corrections"):
                corrected_count += 1

    return json.dumps(
        {
            "results": results,
            "checked": len(results),
            "retracted_count": retracted_count,
            "corrected_count": corrected_count,
        },
        ensure_ascii=False,
    )


@mcp.tool(
    description=(
        "Get papers that cite or are cited by a Zotero item, via OpenAlex. "
        "Use this when the user wants to explore a paper's citation network, "
        "find related work, or trace the influence of a paper. Each result is "
        "flagged with in_library (true/false). Direction: 'cited_by', 'references', or 'both'."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def get_citation_graph(
    item_key: str, direction: str = "both", limit: str | int = 20
) -> str:
    """Get citing and/or referenced works for a Zotero item.

    Args:
        item_key: Zotero item key.
        direction: "cited_by", "references", or "both".
        limit: Max results per direction.

    Returns:
        JSON with cited_by and/or references lists, each with in_library flag.
    """
    from zotero_mcp.openalex_client import OpenAlexClient

    _validate_key(item_key, "item_key")
    limit_int = _clamp_limit(limit, lo=1, hi=50)

    web = _get_web()
    item = web.get_item(item_key.strip())
    if isinstance(item, str):
        return json.dumps({"error": "Could not read item metadata"})

    doi = item.get("DOI", "")
    if not doi:
        return json.dumps(
            {
                "item_key": item_key,
                "error": "No DOI on this item — cannot query citation graph",
            }
        )

    openalex = OpenAlexClient()

    def _add_library_flags(works: list[dict]) -> list[dict]:
        """Batch-check library membership using ThreadPoolExecutor."""
        from concurrent.futures import ThreadPoolExecutor

        dois = [w.get("doi", "") for w in works]

        def _check_doi(d: str) -> dict | None:
            return web._check_duplicate_doi(d) if d else None

        with ThreadPoolExecutor(max_workers=5) as pool:
            existing_items = list(pool.map(_check_doi, dois))

        for work, existing in zip(works, existing_items):
            if existing:
                work["in_library"] = True
                work["zotero_key"] = existing["key"]
            else:
                work["in_library"] = False
        return works

    result: dict = {
        "item_key": item_key,
        "doi": doi,
        "title": item.get("title", ""),
    }

    if direction in ("cited_by", "both"):
        cited_by = openalex.get_citing_works(doi, limit_int)
        result["cited_by"] = _add_library_flags(cited_by)
        result["cited_by_count"] = len(cited_by)

    if direction in ("references", "both"):
        references = openalex.get_references(doi)
        result["references"] = _add_library_flags(references)

    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "List all items in a specific Zotero collection by its key. Use this when "
        "the user wants to see what's in a particular folder/collection."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def get_collection_items(collection_key: str, limit: str | int = 100) -> str:
    """Get items within a collection by its key."""
    _validate_key(collection_key, "collection_key")
    results = _read_local_or_web(
        "get_collection_items", collection_key.strip(), _clamp_limit(limit)
    )
    return json.dumps(results, ensure_ascii=False)


# -- Write tools (web API) --


@mcp.tool(
    description=(
        "Add a paper to Zotero from any identifier: DOI, PMID, or URL (PubMed, "
        "bioRxiv, arXiv, publisher pages). Resolves metadata automatically and "
        "checks for duplicates. Use this when the user wants to save a paper to "
        "their library. Optional title is only used for bare URLs that can't be scraped."
    )
)
@_handle_tool_errors
def create_item(
    input: str,
    title: str | None = None,
    collection_keys: str | list[str] | None = None,
    tags: str | list[str] | None = None,
) -> str:
    """Create a Zotero item from a DOI, PMID, or URL."""
    if not input or not input.strip():
        raise ValueError("input must not be empty")
    input = input.strip()
    collection_keys = _parse_list_param(collection_keys)
    tags = _parse_list_param(tags)
    web = _get_web()
    if input.startswith(("http://", "https://")):
        result = web.create_item_from_url(input, title, collection_keys, tags)
    else:
        result = web.create_item_from_identifier(input, collection_keys, tags)
    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Create a Zotero item with manually provided metadata fields. Use this "
        "instead of create_item when you have structured metadata already (e.g. from "
        "a conversation) rather than a DOI/URL to resolve. Checks for duplicates."
    ),
)
@_handle_tool_errors
def create_item_manual(
    item_type: str,
    title: str,
    creators: str | list[dict] | None = None,
    date: str = "",
    url: str = "",
    doi: str = "",
    publication_title: str = "",
    volume: str = "",
    issue: str = "",
    pages: str = "",
    publisher: str = "",
    abstract: str = "",
    extra: str = "",
    collection_keys: str | list[str] | None = None,
    tags: str | list[str] | None = None,
) -> str:
    """Create a Zotero item with manual metadata."""
    creators = _parse_list_param(creators)
    collection_keys = _parse_list_param(collection_keys)
    tags = _parse_list_param(tags)
    result = _get_web().create_item_manual(
        item_type=item_type,
        title=title,
        creators=creators,
        date=date,
        url=url,
        doi=doi,
        publication_title=publication_title,
        volume=volume,
        issue=issue,
        pages=pages,
        publisher=publisher,
        abstract=abstract,
        extra=extra,
        collection_keys=collection_keys,
        tags=tags,
    )
    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Create a note attached to a Zotero item. Use this to save reading notes, "
        "summaries, or annotations on a paper. Supports HTML or plain text content."
    ),
)
@_handle_tool_errors
def create_note(
    parent_key: str,
    content: str,
    tags: str | list[str] | None = None,
) -> str:
    """Create a child note on a Zotero item."""
    _validate_key(parent_key, "parent_key")
    tags = _parse_list_param(tags)
    result = _get_web().create_note(parent_key.strip(), content, tags)
    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Add tags and/or move multiple items to a collection in one operation. "
        "Use this for bulk organization — e.g. tagging a set of search results or "
        "grouping papers into a collection. Handles rate limiting and version conflicts."
    ),
)
@_handle_tool_errors
def batch_organize(
    item_keys: str | list[str],
    tags: str | list[str] | None = None,
    collection_key: str | None = None,
) -> str:
    """Bulk-add tags and/or collection to multiple items."""
    item_keys = _parse_list_param(item_keys) or []
    tags = _parse_list_param(tags)
    result = _get_web().batch_organize(item_keys, tags, collection_key)
    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Scan the Zotero library for duplicate items using DOI match and title "
        "similarity (>85%). Use this when the user wants to clean up their library "
        "or after bulk imports. Optionally scoped to a single collection."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def find_duplicates(collection_key: str | None = None, limit: str | int = 100) -> str:
    """Find duplicate items in the library or a collection."""
    limit_int = _clamp_limit(limit)
    if collection_key:
        _validate_key(collection_key, "collection_key")
        collection_key = collection_key.strip()
    result = _get_web().find_duplicates(collection_key, limit_int)
    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Create a new collection (folder) in Zotero, optionally nested under a parent. "
        "Use this when the user wants to organize papers into a new group."
    )
)
@_handle_tool_errors
def create_collection(name: str, parent_key: str | None = None) -> str:
    """Create a collection. Returns the new collection key."""
    result = _get_web().create_collection(name, parent_key)
    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Add an existing Zotero item to a collection. Use this to organize a paper "
        "into a folder without moving it from other collections."
    ),
)
@_handle_tool_errors
def add_to_collection(item_key: str, collection_key: str) -> str:
    """Add an existing item to a collection."""
    _validate_key(item_key, "item_key")
    _validate_key(collection_key, "collection_key")
    result = _get_web().add_to_collection(item_key.strip(), collection_key.strip())
    return json.dumps(result, ensure_ascii=False)


_ALLOWED_UPDATE_FIELDS = {
    "title",
    "creators",
    "date",
    "DOI",
    "url",
    "abstractNote",
    "publicationTitle",
    "volume",
    "issue",
    "pages",
    "publisher",
    "ISBN",
    "ISSN",
    "extra",
    "tags",
    "collections",
    "itemType",
    "bookTitle",
    "proceedingsTitle",
    "series",
    "seriesTitle",
    "language",
    "rights",
    "shortTitle",
    "accessDate",
    "archive",
    "archiveLocation",
    "callNumber",
    "libraryCatalog",
    "place",
    "numPages",
    "edition",
    "numberOfVolumes",
}


@mcp.tool(
    description=(
        "Update metadata fields on an existing Zotero item. Use this to correct "
        "titles, authors, dates, DOIs, or other bibliographic fields. "
        "Uses optimistic locking to prevent overwriting concurrent changes."
    ),
)
@_handle_tool_errors
def update_item(item_key: str, fields: dict) -> str:
    """Update item fields. Uses optimistic locking with version check."""
    _validate_key(item_key, "item_key")
    disallowed = set(fields.keys()) - _ALLOWED_UPDATE_FIELDS
    if disallowed:
        raise ValueError(
            f"Fields not allowed for update: {', '.join(sorted(disallowed))}. "
            f"Allowed fields: {', '.join(sorted(_ALLOWED_UPDATE_FIELDS))}"
        )
    result = _get_web().update_item(item_key.strip(), fields)
    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Move Zotero items to trash (reversible). Use this when the user wants to "
        "delete papers. Accepts one or more item keys. Items can be restored from "
        "trash in Zotero. Confirm with user before trashing."
    ),
)
@_handle_tool_errors
def trash_items(item_keys: str | list[str]) -> str:
    """Move items to Zotero trash."""
    keys = _parse_list_param(item_keys) or []
    if not keys:
        raise ValueError("item_keys must not be empty")
    for k in keys:
        _validate_key(k, "item_key")
    result = _get_web().trash_items([k.strip() for k in keys])
    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Permanently delete ALL items in the Zotero trash. THIS IS IRREVERSIBLE. "
        "Always confirm with the user before calling this tool."
    ),
    annotations={"destructiveHint": True},
)
@_handle_tool_errors
def empty_trash() -> str:
    """Permanently delete all trashed items."""
    result = _get_web().empty_trash()
    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Manage tags in your Zotero library. Use this when the user asks about tags, "
        "wants to list/filter tags, remove a tag from all items, or rename a tag. "
        "Actions: 'list' (browse tags, optional prefix filter), "
        "'remove' (delete a tag from every item — requires tag), "
        "'rename' (change a tag name library-wide — requires tag and new_tag)."
    ),
)
@_handle_tool_errors
def manage_tags(
    action: str = "list",
    tag: str = "",
    new_tag: str = "",
    prefix: str = "",
) -> str:
    """Manage tags: list, remove, or rename.

    Args:
        action: One of 'list', 'remove', 'rename'.
        tag: Tag name for remove/rename actions.
        new_tag: New tag name for rename action.
        prefix: Filter prefix for list action.
    """
    action = action.strip().lower()
    web = _get_web()

    if action == "list":
        result = web.get_tags(prefix=prefix or "")
    elif action == "remove":
        if not tag.strip():
            raise ValueError("'remove' action requires tag parameter")
        result = web.remove_tag(tag.strip())
    elif action == "rename":
        if not tag.strip() or not new_tag.strip():
            raise ValueError("'rename' action requires both tag and new_tag parameters")
        result = web.rename_tag(tag.strip(), new_tag.strip())
    else:
        raise ValueError(
            f"Unknown action: {action!r}. Must be: list, remove, rename"
        )
    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Check if preprints have been formally published in a peer-reviewed journal. "
        "Use this when the user has bioRxiv/medRxiv/arXiv papers and wants to know if "
        "a final journal version exists. Reports published DOI, journal name, and "
        "whether the published version is already in the library. Uses CrossRef and OpenAlex."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def check_published_versions(item_keys: str | list[str]) -> str:
    """Check for published journal versions of preprints.

    Args:
        item_keys: Single key or list of Zotero item keys (typically preprints).

    Returns:
        JSON with per-item results and a summary count of items with published versions.
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed

    from zotero_mcp.openalex_client import OpenAlexClient

    keys = _parse_list_param(item_keys) or []
    if not keys:
        raise ValueError("item_keys must not be empty")
    for k in keys:
        _validate_key(k, "item_key")

    web = _get_web()
    openalex = OpenAlexClient()

    published_count = 0
    results: list[dict] = []

    def _check_one(key: str) -> dict:
        item = web.get_item(key.strip())
        if isinstance(item, str):
            return {"key": key, "error": "Could not read item"}

        doi = item.get("DOI", "")
        title = item.get("title", "")
        entry: dict = {"key": key, "title": title, "has_published_version": False}

        if not doi:
            entry["warning"] = "No DOI — cannot check for published version"
            return entry

        entry["doi"] = doi

        # CrossRef is authoritative for preprint→article links
        crossref = web.check_crossref_published(doi)
        published_doi = crossref.get("published_doi")

        # OpenAlex for confirmation and journal name
        oa = openalex.check_published_version(doi)
        entry["is_preprint"] = oa.get("is_preprint", _is_preprint_doi(doi))

        # Use CrossRef DOI if available; fall back to OpenAlex
        if not published_doi and oa.get("published_doi"):
            published_doi = oa["published_doi"]

        if published_doi:
            entry["has_published_version"] = True
            entry["published_doi"] = published_doi
            if oa.get("journal"):
                entry["journal"] = oa["journal"]
            existing = web._check_duplicate_doi(published_doi)
            if existing:
                entry["in_library"] = True
                entry["zotero_key"] = existing.get("key")
            else:
                entry["in_library"] = False

        return entry

    with ThreadPoolExecutor(max_workers=3) as pool:
        futures = {pool.submit(_check_one, k): k for k in keys}
        for future in as_completed(futures):
            entry = future.result()
            results.append(entry)
            if entry.get("has_published_version"):
                published_count += 1

    return json.dumps(
        {
            "results": results,
            "checked": len(results),
            "published_count": published_count,
        },
        ensure_ascii=False,
    )


@mcp.tool(
    description=(
        "Attach a PDF to a Zotero item. Can auto-download a free PDF via "
        "Unpaywall/PMC/bioRxiv, or accept a local file path. Use this when "
        "the user wants to add a PDF to a paper that doesn't have one."
    ),
)
@_handle_tool_errors
def attach_pdf(
    parent_key: str,
    pdf_path: str | None = None,
    doi: str | None = None,
) -> str:
    """Attach a PDF to a Zotero item.

    Args:
        parent_key: Zotero item key to attach the PDF to.
        pdf_path: Local file path to a PDF. If None, tries auto-download.
        doi: DOI to search for free PDF. If None, reads from the item.

    Returns:
        JSON with status, attachment_key, filename, source.
    """
    _validate_key(parent_key, "parent_key")
    if pdf_path:
        if not pdf_path.lower().endswith(".pdf"):
            raise ValueError("pdf_path must be a .pdf file")
        pdf_path = _validate_path(pdf_path, "pdf_path")
    result = _get_web().attach_pdf(parent_key.strip(), pdf_path, doi)
    return json.dumps(result, ensure_ascii=False)


# -- Document tools --


def _fetch_item_metadata(item_keys: list[str]) -> tuple[dict[str, dict], list[str]]:
    """Fetch metadata for multiple items in parallel (local fast path, web fallback).

    Returns:
        Tuple of (item_data dict, missing_keys list).
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed

    def _get_one(key: str) -> dict | str:
        return _read_local_or_web("get_item", key)

    item_data: dict[str, dict] = {}
    missing_keys: list[str] = []

    with ThreadPoolExecutor(max_workers=3) as pool:
        futures = {pool.submit(_get_one, key): key for key in item_keys}
        for future in as_completed(futures):
            key = futures[future]
            try:
                data = future.result()
                if isinstance(data, dict):
                    item_data[key] = data
                else:
                    missing_keys.append(key)
            except Exception:
                missing_keys.append(key)

    return item_data, missing_keys


@mcp.tool(
    description=(
        "Insert live Zotero citation field codes into an existing Word document. "
        "Finds [@ITEM_KEY] markers in the .docx and replaces them with Zotero "
        "field codes. Use this to add citations to a document the user already has. "
        "Preserves existing formatting, styles, images, and layout."
    ),
)
@_handle_tool_errors
def insert_citations(document_path: str, output_path: str | None = None) -> str:
    """Insert Zotero citation field codes into an existing Word document.

    Args:
        document_path: Path to existing .docx with [@ITEM_KEY] markers.
        output_path: Where to save. If omitted, overwrites the original.

    Returns:
        JSON with output_path and citation_count.
    """
    if not document_path.lower().endswith(".docx"):
        raise ValueError("document_path must be a .docx file")
    if output_path and not output_path.lower().endswith(".docx"):
        raise ValueError("output_path must be a .docx file")
    document_path = _validate_path(document_path, "document_path")
    if output_path:
        output_path = _validate_path(output_path, "output_path")

    from zotero_mcp.citation_writer import insert_citations as _insert_citations
    from zotero_mcp.citation_writer import parse_citations

    # Read the document to find all citation keys
    from docx import Document as _Document

    doc = _Document(document_path)
    all_text: list[str] = []
    for para in doc.paragraphs:
        t = "".join(run.text for run in para.runs)
        if t:
            all_text.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = "".join(run.text for run in para.runs)
                    if t:
                        all_text.append(t)

    combined = "\n\n".join(all_text)
    _, key_to_number = parse_citations(combined)
    item_keys = list(key_to_number.keys())

    if not item_keys:
        return json.dumps(
            {
                "output_path": document_path,
                "citation_count": 0,
                "message": "No [@KEY] citation markers found in the document.",
            }
        )

    # Fetch metadata for each item (local fast path, web fallback)
    item_data, missing_keys = _fetch_item_metadata(item_keys)

    user_id = get_config().zotero_user_id or "0"

    result_path, citation_count = _insert_citations(
        document_path, item_data, user_id, output_path
    )

    result = {
        "output_path": result_path,
        "citation_count": citation_count,
    }
    if missing_keys:
        result["missing_keys"] = missing_keys
        result["warning"] = (
            f"Could not fetch metadata for {len(missing_keys)} item(s): "
            f"{', '.join(missing_keys)}. These citations were skipped."
        )
    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Create a new Word document from markdown text with live Zotero citations. "
        "Use [@ITEM_KEY] markers in the content for citations. Use this when writing "
        "a new document from scratch (e.g. literature review, manuscript draft). "
        "For adding citations to an existing document, use insert_citations instead."
    )
)
@_handle_tool_errors
def write_cited_document(content: str, output_path: str) -> str:
    """Write a Word document with live Zotero field codes.

    Args:
        content: Markdown text with [@ITEM_KEY] citation markers.
        output_path: Where to save the .docx file.

    Returns:
        JSON with output_path and citation_count.
    """
    if not output_path.lower().endswith(".docx"):
        raise ValueError("output_path must be a .docx file")
    output_path = _validate_path(output_path, "output_path")

    from zotero_mcp.citation_writer import build_document, parse_citations

    # Extract all unique item keys from the content
    _, key_to_number = parse_citations(content)
    item_keys = list(key_to_number.keys())

    if not item_keys:
        # No citations — just build a plain document
        result_path = build_document(content, {}, "", output_path)
        return json.dumps({"output_path": result_path, "citation_count": 0})

    # Fetch metadata (local fast path, web fallback)
    item_data, missing_keys = _fetch_item_metadata(item_keys)

    user_id = get_config().zotero_user_id or "0"
    result_path = build_document(content, item_data, user_id, output_path)

    result = {
        "output_path": result_path,
        "citation_count": len(item_data),
    }
    if missing_keys:
        result["missing_keys"] = missing_keys
        result["warning"] = (
            f"Could not fetch metadata for {len(missing_keys)} item(s): "
            f"{', '.join(missing_keys)}. These citations were skipped."
        )
    return json.dumps(result, ensure_ascii=False)


# -- Knowledge graph tools --

# Module-level graph cache (avoids rebuilding from SQLite on every query)
_kg_cache: KnowledgeGraph | None = None
_kg_cache_lock = threading.Lock()


def _get_or_build_kg() -> KnowledgeGraph:
    """Return cached KnowledgeGraph, rebuilding from SQLite if needed."""
    global _kg_cache
    with _kg_cache_lock:
        if _kg_cache is not None:
            return _kg_cache
        from zotero_mcp.graph_store import GraphStore
        from zotero_mcp.knowledge_graph import KnowledgeGraph

        store = GraphStore()
        if store.get_last_sync() is None:
            raise RuntimeError(
                "Knowledge graph not yet built. Run build_knowledge_graph first."
            )
        kg = KnowledgeGraph()
        kg.build_from_store(store)
        _kg_cache = kg
        return kg


def _invalidate_kg_cache() -> None:
    """Invalidate the cached graph after build/sync."""
    global _kg_cache
    with _kg_cache_lock:
        _kg_cache = None


def _extract_openalex_id(url: str) -> str:
    """Extract OpenAlex ID from URL like 'https://openalex.org/W123'."""
    return url.split("/")[-1] if "/" in url else url


def _index_works(works, key_by_doi, store, openalex):
    """Index OpenAlex works into the graph store.

    Shared logic for both full build and incremental sync: stores papers,
    topics, authors, resolves references, and stores citation edges.

    Returns:
        Dict with papers_indexed, citations_indexed, references_resolved,
        topics_indexed, authors_indexed counts.
    """
    from zotero_mcp.openalex_client import OpenAlexClient

    papers_added = 0
    topics_indexed = 0
    authors_indexed = 0
    all_ref_ids: set[str] = set()
    work_refs: dict[str, list[str]] = {}

    for work in works:
        doi = (work.get("doi") or "").replace("https://doi.org/", "")
        if not doi:
            continue
        authorships = work.get("authorships", [])
        authors = "; ".join(
            a.get("author", {}).get("display_name", "") for a in authorships[:3]
        )
        pub_date = (work.get("publication_date") or "")[:7]  # YYYY-MM
        abstract = OpenAlexClient.reconstruct_abstract(work)
        store.upsert_paper(
            doi=doi,
            zotero_key=key_by_doi.get(doi, ""),
            title=work.get("title", ""),
            year=work.get("publication_year", 0),
            authors=authors,
            openalex_id=work.get("id", ""),
            publication_date=pub_date,
            abstract=abstract,
        )
        papers_added += 1

        for topic in OpenAlexClient.extract_topics(work):
            store.upsert_topic(doi=doi, **topic)
            topics_indexed += 1

        for auth in OpenAlexClient.extract_authorships(work):
            store.upsert_author(
                openalex_author_id=auth["openalex_author_id"],
                display_name=auth["display_name"],
                orcid=auth.get("orcid", ""),
                institution=auth.get("institution", ""),
            )
            store.upsert_paper_author(
                doi=doi,
                openalex_author_id=auth["openalex_author_id"],
                position=auth.get("position", 0),
            )
            authors_indexed += 1

        ref_ids = [
            _extract_openalex_id(url) for url in work.get("referenced_works", [])
        ]
        if ref_ids:
            work_refs[doi] = ref_ids
            all_ref_ids.update(ref_ids)

    # Resolve OpenAlex IDs to DOIs
    known_oa_ids = {_extract_openalex_id(w.get("id", "")) for w in works if w.get("id")}
    unknown_ids = list(all_ref_ids - known_oa_ids)

    id_to_doi: dict[str, str] = {}
    for w in works:
        oa_id = _extract_openalex_id(w.get("id", ""))
        d = (w.get("doi") or "").replace("https://doi.org/", "")
        if oa_id and d:
            id_to_doi[oa_id] = d

    if unknown_ids:
        resolved = openalex.resolve_ids_to_dois(unknown_ids)
        id_to_doi.update(resolved)
        for oa_id, ref_doi in resolved.items():
            store.upsert_paper(
                doi=ref_doi,
                zotero_key="",
                title="",
                year=0,
                authors="",
                openalex_id=f"https://openalex.org/{oa_id}",
            )

    citations_added = 0
    for citing_doi, ref_ids in work_refs.items():
        for ref_id in ref_ids:
            cited_doi = id_to_doi.get(ref_id)
            if cited_doi:
                store.upsert_citation(citing_doi=citing_doi, cited_doi=cited_doi)
                citations_added += 1

    return {
        "papers_indexed": papers_added,
        "citations_indexed": citations_added,
        "references_resolved": len(id_to_doi),
        "topics_indexed": topics_indexed,
        "authors_indexed": authors_indexed,
    }


def _build_knowledge_graph(full_rebuild: bool = False) -> dict:
    """Build or incrementally update the knowledge graph. Returns stats dict."""
    from datetime import datetime, timezone

    from zotero_mcp.graph_store import GraphStore
    from zotero_mcp.openalex_client import OpenAlexClient

    web = _get_web()
    openalex = OpenAlexClient()
    store = GraphStore()

    last_sync = store.get_last_sync()
    is_incremental = last_sync is not None and not full_rebuild

    items = web.get_all_items_with_dois()
    if not items:
        return {"error": "No items with DOIs found in library"}

    if is_incremental:
        existing_dois = store.get_doi_set()
        items = [item for item in items if item["DOI"] not in existing_dois]
        if not items:
            kg = _get_or_build_kg()
            stats = kg.get_stats()
            stats["new_papers"] = 0
            stats["new_citations"] = 0
            stats["mode"] = "sync"
            return stats

    doi_list = [item["DOI"] for item in items]
    key_by_doi = {item["DOI"]: item["key"] for item in items}

    works = openalex.bulk_get_works(doi_list)
    counts = _index_works(works, key_by_doi, store, openalex)

    store.set_last_sync(datetime.now(timezone.utc).isoformat())
    _invalidate_kg_cache()
    kg = _get_or_build_kg()
    stats = kg.get_stats()
    stats["mode"] = "sync" if is_incremental else "full_build"

    if is_incremental:
        stats["new_papers"] = counts["papers_indexed"]
        stats["new_citations"] = counts["citations_indexed"]
        stats["new_topics"] = counts["topics_indexed"]
        stats["new_authors"] = counts["authors_indexed"]
    else:
        stats.update(counts)

    return stats


def _build_fulltext_index(full_rebuild: bool = False, limit: int = 0) -> dict:
    """Build or update the full-text search index. Returns stats dict."""
    from concurrent.futures import ThreadPoolExecutor, as_completed

    from zotero_mcp.graph_store import GraphStore
    from zotero_mcp.text_extractor import extract_text_from_pdf, index_paper_text

    web = _get_web()
    store = GraphStore()

    items = web.get_all_items_with_dois()
    if not items:
        return {"error": "No items with DOIs found in library"}

    already_indexed = store.get_indexed_dois()
    total = len(items)

    if not full_rebuild:
        items = [it for it in items if it["DOI"] not in already_indexed]
    skipped = total - len(items)

    if limit > 0:
        items = items[:limit]

    indexed = 0
    failed = 0

    def _process_one(item: dict) -> dict:
        """Extract and index text for a single item."""
        item_key = item["key"]
        item_doi = item["DOI"]

        try:
            children = _read_local_or_web(
                "get_children", item_key, item_type="attachment"
            )
        except Exception:
            return {"doi": item_doi, "status": "failed", "reason": "no_attachments"}

        pdf_atts = [
            c for c in children if c.get("contentType") == "application/pdf"
        ]
        if not pdf_atts:
            return {"doi": item_doi, "status": "failed", "reason": "no_pdf"}

        att = pdf_atts[0]
        att_key = att.get("key", "")
        pdf_source = None

        try:
            local = _get_local()
            local_path = local.get_attachment_path(att_key)
            if local_path:
                pdf_source = local_path
        except Exception:
            pass

        if pdf_source is None:
            try:
                pdf_source = web.download_attachment(att_key)
            except Exception:
                return {"doi": item_doi, "status": "failed", "reason": "download_failed"}

        text = extract_text_from_pdf(pdf_source)
        if not text:
            return {"doi": item_doi, "status": "failed", "reason": "no_text_extracted"}

        index_paper_text(store, item_doi, text)
        return {"doi": item_doi, "status": "indexed"}

    with ThreadPoolExecutor(max_workers=4) as pool:
        futures = {pool.submit(_process_one, it): it for it in items}
        for future in as_completed(futures):
            try:
                result = future.result()
                if result["status"] == "indexed":
                    indexed += 1
                else:
                    failed += 1
            except Exception:
                failed += 1

    return {
        "indexed": indexed,
        "skipped": skipped,
        "failed": failed,
        "total": total,
    }


@mcp.tool(
    description=(
        "Build or update indexes for your Zotero library. Use this after adding "
        "new papers to enable graph queries and full-text search. "
        "Types: 'graph' (citation network + analytics via OpenAlex — enables "
        "query_knowledge_graph, query_authors, export_knowledge_graph), "
        "'fulltext' (PDF text extraction + FTS5 index — enables search_fulltext), "
        "'both' (runs graph then fulltext). "
        "Auto-detects full build vs incremental sync. "
        "Set full_rebuild=true to force a complete rebuild."
    ),
)
@_handle_tool_errors
def build_index(
    type: str = "graph",
    full_rebuild: bool = False,
    limit: int = 0,
) -> str:
    """Build or update knowledge graph and/or fulltext index.

    Args:
        type: One of 'graph', 'fulltext', 'both'.
        full_rebuild: If True, force complete rebuild instead of incremental.
        limit: For fulltext, cap number of papers to process (0 = all).
    """
    type = type.strip().lower()
    if type not in ("graph", "fulltext", "both"):
        raise ValueError(
            f"Unknown type: {type!r}. Must be: graph, fulltext, both"
        )

    result: dict = {}
    if type in ("graph", "both"):
        result["graph"] = _build_knowledge_graph(full_rebuild)
    if type in ("fulltext", "both"):
        result["fulltext"] = _build_fulltext_index(full_rebuild, limit)

    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Query the knowledge graph for insights about your library's citation network. "
        "Use this when the user asks about influential papers, research clusters, "
        "publication trends, or relationships between papers. "
        "Query types: 'influential' (PageRank-ranked papers), "
        "'clusters' (research topic groupings), "
        "'bridges' (papers connecting different clusters), "
        "'path' (shortest citation path between two DOIs — requires doi_a and doi_b), "
        "'neighborhood' (papers within N hops of a DOI — requires doi and optional depth), "
        "'stats' (graph summary), "
        "'timeline' (papers per month — optional topic filter, start_year, end_year), "
        "'topic_evolution' (per-subfield paper counts by month — optional start_year, end_year), "
        "'citation_velocity' (month-by-month citation count for a DOI — requires doi), "
        "'trending' (papers with accelerating citation rates — optional limit, years window). "
        "Requires build_index(type='graph') to be run first."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def query_knowledge_graph(
    query_type: str,
    doi: str = "",
    doi_a: str = "",
    doi_b: str = "",
    depth: int = 1,
    limit: int = 10,
    topic: str = "",
    start_year: int = 0,
    end_year: int = 0,
    years: int = 3,
) -> str:
    """Query the knowledge graph (uses cached graph)."""
    kg = _get_or_build_kg()

    if query_type == "influential":
        result = kg.get_influential_papers(top_n=limit)
    elif query_type == "clusters":
        result = kg.get_clusters()
    elif query_type == "bridges":
        result = kg.get_bridge_papers(top_n=limit)
    elif query_type == "path":
        if not doi_a or not doi_b:
            raise ValueError("path query requires doi_a and doi_b")
        result = kg.get_path(doi_a, doi_b)
    elif query_type == "neighborhood":
        if not doi:
            raise ValueError("neighborhood query requires doi")
        result = kg.get_neighborhood(doi, depth=depth)
    elif query_type == "stats":
        result = kg.get_stats()
    elif query_type == "timeline":
        result = kg.get_timeline(
            topic=topic or None,
            start_year=start_year or None,
            end_year=end_year or None,
        )
    elif query_type == "topic_evolution":
        result = kg.get_topic_evolution(
            start_year=start_year or None,
            end_year=end_year or None,
            limit=limit,
        )
    elif query_type == "citation_velocity":
        if not doi:
            raise ValueError("citation_velocity query requires doi")
        result = kg.get_citation_velocity(doi)
    elif query_type == "trending":
        result = kg.get_trending(top_n=limit, years=years)
    else:
        raise ValueError(
            f"Unknown query_type: {query_type!r}. "
            "Must be: influential, clusters, bridges, path, neighborhood, "
            "stats, timeline, topic_evolution, citation_velocity, trending"
        )

    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Find papers related to items in your library via Semantic Scholar "
        "recommendations (similar to Connected Papers or ResearchRabbit). "
        "Use this when the user wants to discover new papers on a topic. "
        "Provide one or more item keys as seeds — the more seeds, the better "
        "the recommendations. Each result is flagged with in_library (true/false)."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def find_related_papers(
    item_keys: str | list[str],
    limit: str | int = 10,
) -> str:
    """Get paper recommendations from Semantic Scholar."""
    from zotero_mcp.semantic_scholar_client import SemanticScholarClient

    keys = _parse_list_param(item_keys) or []
    if not keys:
        raise ValueError("item_keys must not be empty")
    for k in keys:
        _validate_key(k, "item_key")

    limit_int = _clamp_limit(limit, lo=1, hi=50)
    web = _get_web()

    # Resolve keys to DOIs
    dois: list[str] = []
    for key in keys:
        item = web.get_item(key.strip())
        if isinstance(item, dict) and item.get("DOI"):
            dois.append(item["DOI"])

    if not dois:
        return json.dumps(
            {
                "error": "None of the provided items have DOIs",
                "item_keys": keys,
            }
        )

    s2 = SemanticScholarClient(api_key=get_config().semantic_scholar_api_key)
    recommendations = s2.get_recommendations(dois, limit=limit_int)

    # Flag which recommendations are already in library
    for rec in recommendations:
        rec_doi = rec.get("doi", "")
        if rec_doi:
            existing = web._check_duplicate_doi(rec_doi)
            if existing:
                rec["in_library"] = True
                rec["zotero_key"] = existing["key"]
            else:
                rec["in_library"] = False
        else:
            rec["in_library"] = False

    return json.dumps(
        {
            "seed_count": len(dois),
            "recommendations": recommendations,
        },
        ensure_ascii=False,
    )


# -- Full-text search tools --


@mcp.tool(
    description=(
        "Search the full text of indexed PDFs in your library. Use this when "
        "the user wants to find papers that mention a specific term, method, "
        "drug, or concept in their body text (not just titles/abstracts). "
        "Returns matching papers with highlighted text snippets. "
        "Requires build_index(type='fulltext') to be run first."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def search_fulltext(query: str, limit: str | int = 20) -> str:
    """Search full-text index for matching papers.

    Args:
        query: FTS5 search query string.
        limit: Maximum results to return.

    Returns:
        JSON with matching papers and highlighted snippets.
    """
    from zotero_mcp.graph_store import GraphStore

    if not query or not query.strip():
        raise ValueError("query must not be empty")

    limit_int = _clamp_limit(limit, lo=1, hi=100)
    store = GraphStore()
    results = store.search_fulltext(query.strip(), limit_int)
    return json.dumps(
        {"query": query.strip(), "results": results, "count": len(results)},
        ensure_ascii=False,
    )


@mcp.tool(
    description=(
        "Query the author co-authorship network in your knowledge graph. "
        "Use this when the user asks about who publishes most, who collaborates "
        "with whom, or wants to map out an author's network. "
        "Query types: 'prolific' (authors by paper count), "
        "'influential' (authors by summed PageRank of their papers), "
        "'coauthors_of' (co-authors of a named author, ranked by shared papers), "
        "'network' (ego network for an author within N hops — requires author_name, optional depth), "
        "'clusters' (author community groupings). "
        "Requires build_index(type='graph') to be run first."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def query_authors(
    query_type: str,
    author_name: str = "",
    limit: int = 10,
    depth: int = 1,
) -> str:
    """Query the author co-citation network."""
    kg = _get_or_build_kg()

    if query_type == "prolific":
        result = kg.get_prolific_authors(top_n=limit)
    elif query_type == "influential":
        result = kg.get_influential_authors(top_n=limit)
    elif query_type == "coauthors_of":
        if not author_name:
            raise ValueError("coauthors_of query requires author_name")
        author_id = kg._resolve_author(author_name)
        result = kg.get_coauthors_of(author_id, top_n=limit)
    elif query_type == "network":
        if not author_name:
            raise ValueError("network query requires author_name")
        author_id = kg._resolve_author(author_name)
        result = kg.get_author_network(author_id, depth=depth)
    elif query_type == "clusters":
        result = kg.get_author_clusters()
    else:
        raise ValueError(
            f"Unknown query_type: {query_type!r}. "
            "Must be: prolific, influential, coauthors_of, network, clusters"
        )

    return json.dumps(result, ensure_ascii=False)


@mcp.tool(
    description=(
        "Export the knowledge graph as an interactive HTML visualization that opens "
        "in any browser. Use this when the user wants to see their citation network "
        "visually, explore research clusters, or share a graph. "
        "Views: 'citations' (paper nodes + citation edges, colored by cluster), "
        "'authors' (author nodes + co-authorship edges), "
        "'full' (both layers, papers capped at 200 by PageRank). "
        "Requires build_index(type='graph') to have been run first."
    )
)
@_handle_tool_errors
def export_knowledge_graph(
    view: str = "citations",
    path: str = "",
) -> str:
    """Export knowledge graph as interactive HTML.

    Args:
        view: One of 'citations', 'authors', 'full'.
        path: Output file path. If empty, writes to a temp file.

    Returns:
        JSON with path, node/edge counts, and view type.
    """
    from zotero_mcp.graph_renderer import (
        render_authors_view,
        render_citations_view,
        render_full_view,
    )

    kg = _get_or_build_kg()

    if view == "authors":
        html, stats = render_authors_view(kg)
    elif view == "full":
        html, stats = render_full_view(kg)
    else:
        html, stats = render_citations_view(kg)

    if not path:
        fd, path = tempfile.mkstemp(suffix=".html", prefix="zotero_mcp_graph_")
        os.close(fd)
        _register_temp_file(path)

    with open(path, "w", encoding="utf-8") as f:
        f.write(html)

    return json.dumps({"path": path, **stats}, ensure_ascii=False)


# -- Entity extraction tools --


@mcp.tool(
    description=(
        "Get papers that have abstracts but no extracted biomedical entities yet. "
        "Use this to find papers needing entity extraction, then extract entities "
        "(conditions, drugs, genes, biomarkers, methods, outcomes) from the returned "
        "abstracts and save them with store_entities."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def get_unextracted_abstracts(limit: str | int = 50) -> str:
    """Get papers with abstracts that haven't had entities extracted.

    Args:
        limit: Maximum number of papers to return.

    Returns:
        JSON with papers list and remaining count.
    """
    from zotero_mcp.graph_store import GraphStore

    limit_int = _clamp_limit(limit, lo=1, hi=200)
    store = GraphStore()
    all_unextracted = store.get_unextracted_dois()
    batch = all_unextracted[:limit_int]
    remaining = len(all_unextracted) - len(batch)
    return json.dumps(
        {"papers": batch, "remaining": remaining},
        ensure_ascii=False,
    )


@mcp.tool(
    description=(
        "Store extracted biomedical entities for papers. Call this after extracting "
        "entities from abstracts (via get_unextracted_abstracts or any paper reading). "
        "Entity types: condition, biomarker, drug, method, gene, organism, outcome, dataset. "
        "Input: list of {doi, entities: [{name, type}]}."
    ),
)
@_handle_tool_errors
def store_entities(results: str | list) -> str:
    """Store extracted entities for one or more papers.

    Args:
        results: JSON list of {doi, entities: [{name, type}]}.
            May be a JSON string or a parsed list.

    Returns:
        JSON with stored paper count and entity creation stats.
    """
    from zotero_mcp.graph_store import GraphStore

    if isinstance(results, str):
        results = json.loads(results)
    if not isinstance(results, list):
        raise ValueError("results must be a list of {doi, entities: [{name, type}]}")

    store = GraphStore()
    papers_stored = 0
    entities_created = 0
    entities_reused = 0

    for item in results:
        doi = item.get("doi", "").strip()
        if not doi:
            continue
        entities = item.get("entities", [])
        if not entities:
            continue

        entities_stored_for_paper = 0
        for ent in entities:
            ent_name = ent.get("name", "").strip()
            ent_type = ent.get("type", "").strip()
            if not ent_name or not ent_type:
                continue

            already_exists = store.entity_exists(ent_name, ent_type)
            entity_id = store.upsert_entity(ent_name, ent_type)

            if already_exists:
                entities_reused += 1
            else:
                entities_created += 1

            store.upsert_paper_entity(doi, entity_id)
            entities_stored_for_paper += 1

        if entities_stored_for_paper > 0:
            papers_stored += 1

    _invalidate_kg_cache()
    return json.dumps(
        {
            "stored": papers_stored,
            "entities_created": entities_created,
            "entities_reused": entities_reused,
        },
        ensure_ascii=False,
    )


@mcp.tool(
    description=(
        "Search the biomedical entity graph. Use this when the user asks about "
        "which papers mention a condition/drug/gene, what entities co-occur, "
        "or what two papers have in common. "
        "Query types: 'by_name' (papers mentioning an entity), "
        "'by_type' (common entities of a type, or list all types if no type given), "
        "'co_occurrence' (entities that co-occur with a given entity), "
        "'shared_entities' (entities shared by two papers — requires doi_a and doi_b), "
        "'paper_entities' (all entities extracted from a paper — requires doi)."
    ),
    annotations={"readOnlyHint": True},
)
@_handle_tool_errors
def search_entities(
    query_type: str,
    entity_name: str = "",
    entity_type: str = "",
    doi: str = "",
    doi_a: str = "",
    doi_b: str = "",
    limit: str | int = 20,
) -> str:
    """Search the entity graph with various query types.

    Args:
        query_type: One of by_name, by_type, co_occurrence,
            shared_entities, paper_entities.
        entity_name: Entity name for by_name/co_occurrence queries.
        entity_type: Entity type filter for by_type queries.
        doi: Paper DOI for paper_entities query.
        doi_a: First DOI for shared_entities query.
        doi_b: Second DOI for shared_entities query.
        limit: Maximum results.

    Returns:
        JSON with query results.
    """
    from zotero_mcp.graph_store import GraphStore

    limit_int = _clamp_limit(limit, lo=1, hi=100)
    store = GraphStore()

    if query_type == "by_name":
        if not entity_name:
            raise ValueError("by_name query requires entity_name")
        entities = store.search_entities_by_name(entity_name, limit=limit_int)
        results = []
        for ent in entities:
            papers = store.get_papers_for_entity(ent["entity_id"])
            results.append({
                **ent,
                "paper_count": len(papers),
                "papers": [
                    {"doi": p["doi"], "title": p["title"], "year": p["year"]}
                    for p in papers[:limit_int]
                ],
            })
        return json.dumps({"query": "by_name", "results": results}, ensure_ascii=False)

    elif query_type == "by_type":
        if entity_type:
            results = store.get_entities_by_type(entity_type, limit=limit_int)
            return json.dumps(
                {"query": "by_type", "entity_type": entity_type,
                 "results": results},
                ensure_ascii=False,
            )
        else:
            types = store.get_all_entity_types()
            return json.dumps(
                {"query": "by_type", "entity_types": types},
                ensure_ascii=False,
            )

    elif query_type == "co_occurrence":
        if not entity_name:
            raise ValueError("co_occurrence query requires entity_name")
        matches = store.search_entities_by_name(entity_name, limit=1)
        if not matches:
            return json.dumps(
                {"query": "co_occurrence", "error": f"Entity not found: {entity_name}"},
                ensure_ascii=False,
            )
        entity_id = matches[0]["entity_id"]
        co_occurring = store.get_entity_co_occurrence(entity_id, limit=limit_int)
        return json.dumps(
            {"query": "co_occurrence", "entity": matches[0],
             "co_occurring": co_occurring},
            ensure_ascii=False,
        )

    elif query_type == "shared_entities":
        if not doi_a or not doi_b:
            raise ValueError("shared_entities query requires doi_a and doi_b")
        shared = store.get_shared_entities(doi_a, doi_b)
        return json.dumps(
            {"query": "shared_entities", "doi_a": doi_a, "doi_b": doi_b,
             "shared": shared},
            ensure_ascii=False,
        )

    elif query_type == "paper_entities":
        if not doi:
            raise ValueError("paper_entities query requires doi")
        entities = store.get_entities_for_doi(doi)
        return json.dumps(
            {"query": "paper_entities", "doi": doi, "entities": entities},
            ensure_ascii=False,
        )

    else:
        raise ValueError(
            f"Unknown query_type: {query_type!r}. "
            "Must be: by_name, by_type, co_occurrence, shared_entities, paper_entities"
        )


# -- MCP Prompts (multi-tool workflows) --


@mcp.prompt(
    name="literature_audit",
    description=(
        "Run a full literature audit on selected papers: check for retractions, "
        "verify preprint publication status, and scan for duplicates."
    ),
)
def literature_audit(item_keys: str = "") -> str:
    """Guide for running a comprehensive literature audit."""
    return (
        "Run a literature quality audit on the user's Zotero library:\n\n"
        "1. First, call check_retractions with the item keys to identify any "
        "retracted or corrected papers.\n"
        "2. Then call check_published_versions to find preprints that now have "
        "published journal versions.\n"
        "3. Finally, call find_duplicates to identify duplicate entries.\n"
        "4. Summarize the findings: retracted papers (urgent), preprints with "
        "published versions (suggest updating), and duplicates (suggest merging).\n\n"
        f"Item keys to audit: {item_keys or '(search the library first with search_items)'}"
    )


@mcp.prompt(
    name="build_and_explore",
    description=(
        "Build the knowledge graph and fulltext index, then explore the library's "
        "research landscape with influential papers, clusters, and trends."
    ),
)
def build_and_explore() -> str:
    """Guide for building indexes and exploring the library."""
    return (
        "Build the library's knowledge base and explore its research landscape:\n\n"
        "1. Call build_index(type='both') to build the citation graph and "
        "fulltext search index.\n"
        "2. Call query_knowledge_graph(query_type='stats') for an overview.\n"
        "3. Call query_knowledge_graph(query_type='influential') to find the "
        "most influential papers by PageRank.\n"
        "4. Call query_knowledge_graph(query_type='clusters') to discover "
        "research topic groupings.\n"
        "5. Call query_knowledge_graph(query_type='trending') to find papers "
        "with accelerating citation rates.\n"
        "6. Present a structured summary of the library's research landscape."
    )


@mcp.prompt(
    name="add_and_verify",
    description=(
        "Add a paper to the library by DOI/PMID/URL, then verify it: "
        "check for retractions, find related work, and attach a PDF."
    ),
)
def add_and_verify(identifier: str = "") -> str:
    """Guide for adding a paper and running verification checks."""
    return (
        "Add a paper to the library and run verification checks:\n\n"
        f"1. Call create_item(input='{identifier or '<DOI, PMID, or URL>'}') "
        "to add the paper.\n"
        "2. Call check_retractions with the new item key to verify it hasn't "
        "been retracted.\n"
        "3. Call attach_pdf with the item key to try auto-downloading the PDF.\n"
        "4. Call find_related_papers with the item key to discover related work.\n"
        "5. Report the results: paper added, retraction status, PDF status, "
        "and top related papers the user might want to add."
    )


@mcp.prompt(
    name="extract_entities",
    description=(
        "Extract biomedical entities from papers that haven't been processed yet, "
        "then store them for search and co-occurrence analysis."
    ),
)
def extract_entities_prompt() -> str:
    """Guide for the entity extraction workflow."""
    return (
        "Extract biomedical entities from unprocessed papers:\n\n"
        "1. Call get_unextracted_abstracts(limit=20) to get papers needing "
        "entity extraction.\n"
        "2. For each paper's abstract, identify biomedical entities:\n"
        "   - conditions (diseases, symptoms)\n"
        "   - drugs (medications, compounds)\n"
        "   - genes (gene names, variants)\n"
        "   - biomarkers (lab values, molecular markers)\n"
        "   - methods (study designs, techniques)\n"
        "   - outcomes (endpoints, measures)\n"
        "   - organisms (species, model organisms)\n"
        "   - datasets (databases, registries)\n"
        "3. Call store_entities with the extracted entities.\n"
        "4. Report how many papers were processed and entities found.\n"
        "5. If there are remaining unprocessed papers, offer to continue."
    )
