"""Tests for citation_writer -- Zotero field code generation."""

import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from zotero_mcp.citation_writer import (
    TextBlock,
    parse_citations,
    zotero_to_csl_json,
    add_citation_field,
    add_bibliography_field,
    build_document,
    insert_citations,
)

# -- Parser tests --


def test_parse_single_citation():
    blocks, mapping = parse_citations("text [@ABC123] more")
    assert len(blocks) == 3
    assert blocks[0] == TextBlock("text", "text ", [], [])
    assert blocks[1] == TextBlock("citation", "", ["ABC123"], [1])
    assert blocks[2] == TextBlock("text", " more", [], [])
    assert mapping == {"ABC123": 1}


def test_parse_grouped_citations():
    blocks, mapping = parse_citations("text [@AAA, @BBB]")
    assert len(blocks) == 2
    citation = blocks[1]
    assert citation.keys == ["AAA", "BBB"]
    assert citation.numbers == [1, 2]
    assert mapping == {"AAA": 1, "BBB": 2}


def test_parse_duplicate_same_number():
    blocks, mapping = parse_citations("first [@ABC] second [@ABC]")
    assert mapping == {"ABC": 1}
    assert blocks[1].numbers == [1]
    assert blocks[3].numbers == [1]


def test_parse_no_citations():
    blocks, mapping = parse_citations("plain text no citations")
    assert len(blocks) == 1
    assert blocks[0].kind == "text"
    assert mapping == {}


def test_parse_sequential_numbering():
    _, mapping = parse_citations("[@A] then [@B] then [@C]")
    assert mapping == {"A": 1, "B": 2, "C": 3}


# -- CSL-JSON conversion tests --

SAMPLE_ITEM = {
    "key": "ABC123",
    "itemType": "journalArticle",
    "title": "Test Article Title",
    "creators": [
        {"creatorType": "author", "firstName": "John", "lastName": "Doe"},
        {"creatorType": "author", "firstName": "Jane", "lastName": "Smith"},
        {"creatorType": "editor", "firstName": "Ed", "lastName": "Itor"},
    ],
    "date": "2024-03-15",
    "DOI": "10.1234/test",
    "publicationTitle": "Journal of Testing",
    "volume": "42",
    "issue": "3",
    "pages": "100-110",
    "ISSN": "1234-5678",
}


def test_journal_article_conversion():
    csl = zotero_to_csl_json(SAMPLE_ITEM, "12345")
    assert csl["type"] == "article-journal"
    assert csl["title"] == "Test Article Title"
    assert csl["container-title"] == "Journal of Testing"
    assert csl["volume"] == "42"
    assert csl["issue"] == "3"
    assert csl["page"] == "100-110"
    assert csl["DOI"] == "10.1234/test"


def test_creator_mapping():
    csl = zotero_to_csl_json(SAMPLE_ITEM, "12345")
    # Only authors, not editors
    assert len(csl["author"]) == 2
    assert csl["author"][0] == {"family": "Doe", "given": "John"}
    assert csl["author"][1] == {"family": "Smith", "given": "Jane"}


def test_date_parsing_full():
    csl = zotero_to_csl_json({**SAMPLE_ITEM, "date": "2024-03-15"}, "12345")
    assert csl["issued"] == {"date-parts": [[2024, 3, 15]]}


def test_date_parsing_year_only():
    csl = zotero_to_csl_json({**SAMPLE_ITEM, "date": "2024"}, "12345")
    assert csl["issued"] == {"date-parts": [[2024]]}


def test_date_parsing_year_month():
    csl = zotero_to_csl_json({**SAMPLE_ITEM, "date": "2024-03"}, "12345")
    assert csl["issued"] == {"date-parts": [[2024, 3]]}


def test_uri_construction():
    csl = zotero_to_csl_json(SAMPLE_ITEM, "12345")
    assert "http://zotero.org/users/12345/items/ABC123" in csl["_uris"]


def test_unknown_item_type_fallback():
    item = {**SAMPLE_ITEM, "itemType": "podcast"}
    csl = zotero_to_csl_json(item, "12345")
    assert csl["type"] == "article"


# -- Field code tests --


def test_citation_field_code_structure():
    doc = Document()
    para = doc.add_paragraph()
    citation_json = {
        "citationID": "test1",
        "properties": {
            "formattedCitation": "1",
            "plainCitation": "1",
            "noteIndex": 0,
        },
        "citationItems": [
            {
                "id": 1,
                "uris": ["http://example.com"],
                "itemData": {"type": "article-journal", "title": "Test"},
            }
        ],
        "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json",
    }
    add_citation_field(para, citation_json, "1")

    # Check XML has fldChar elements
    xml = para._element.xml
    assert "w:fldChar" in xml
    assert "ADDIN ZOTERO_ITEM CSL_CITATION" in xml


def test_citation_instrtext_valid_json():
    doc = Document()
    para = doc.add_paragraph()
    csl_data = {"type": "article-journal", "title": "Test Paper"}
    citation_json = {
        "citationID": "test1",
        "properties": {
            "formattedCitation": "1",
            "plainCitation": "1",
            "noteIndex": 0,
        },
        "citationItems": [
            {"id": 1, "uris": ["http://example.com"], "itemData": csl_data}
        ],
        "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json",
    }
    add_citation_field(para, citation_json, "1")

    # Extract instrText and verify it's valid JSON after the prefix
    instr_elements = para._element.findall(".//" + qn("w:instrText"))
    assert len(instr_elements) >= 1
    instr_text = instr_elements[0].text
    json_str = instr_text.replace("ADDIN ZOTERO_ITEM CSL_CITATION ", "")
    parsed = json.loads(json_str)
    assert parsed["citationID"] == "test1"


def test_bibliography_field_code():
    doc = Document()
    para = doc.add_paragraph()
    add_bibliography_field(para)
    xml = para._element.xml
    assert "ADDIN ZOTERO_BIBL" in xml
    assert "CSL_BIBLIOGRAPHY" in xml


def test_superscript_display():
    doc = Document()
    para = doc.add_paragraph()
    citation_json = {
        "citationID": "test1",
        "properties": {
            "formattedCitation": "1",
            "plainCitation": "1",
            "noteIndex": 0,
        },
        "citationItems": [
            {
                "id": 1,
                "uris": ["uri"],
                "itemData": {"type": "article-journal", "title": "T"},
            }
        ],
        "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json",
    }
    add_citation_field(para, citation_json, "1")
    xml = para._element.xml
    assert "w:vertAlign" in xml
    assert "superscript" in xml


# -- Document assembly tests --


def test_build_simple_document(tmp_path):
    output = tmp_path / "test.docx"
    item_data = {
        "ABC": {
            "key": "ABC",
            "itemType": "journalArticle",
            "title": "Test Paper",
            "creators": [
                {"creatorType": "author", "firstName": "J", "lastName": "Doe"}
            ],
            "date": "2024",
            "DOI": "10.1/test",
            "publicationTitle": "Test Journal",
            "volume": "1",
            "issue": "1",
            "pages": "1-5",
        }
    }
    result = build_document(
        "This is a test [@ABC].",
        item_data,
        "12345",
        str(output),
    )
    assert Path(result).exists()
    doc = Document(result)
    # Should have at least: content paragraph + References heading + bibliography paragraph
    assert len(doc.paragraphs) >= 3


def test_build_with_headings(tmp_path):
    output = tmp_path / "headings.docx"
    result = build_document(
        "# Introduction\n\nSome text.\n\n## Methods\n\nMore text.",
        {},
        "12345",
        str(output),
    )
    doc = Document(result)
    styles = [p.style.name for p in doc.paragraphs]
    assert "Heading 1" in styles
    assert "Heading 2" in styles


def test_build_multiple_citations_sequential(tmp_path):
    output = tmp_path / "multi.docx"
    item_data = {
        "A": {
            "key": "A",
            "itemType": "journalArticle",
            "title": "Paper A",
            "creators": [],
            "date": "2024",
            "DOI": "",
            "publicationTitle": "J",
            "volume": "",
            "issue": "",
            "pages": "",
        },
        "B": {
            "key": "B",
            "itemType": "journalArticle",
            "title": "Paper B",
            "creators": [],
            "date": "2024",
            "DOI": "",
            "publicationTitle": "J",
            "volume": "",
            "issue": "",
            "pages": "",
        },
    }
    result = build_document("First [@A] then [@B].", item_data, "12345", str(output))
    doc = Document(result)
    xml = doc.paragraphs[0]._element.xml
    # Both citations should be present
    assert "ADDIN ZOTERO_ITEM" in xml


def test_build_grouped_citation_display(tmp_path):
    output = tmp_path / "grouped.docx"
    item_data = {
        "A": {
            "key": "A",
            "itemType": "journalArticle",
            "title": "Paper A",
            "creators": [],
            "date": "2024",
            "DOI": "",
            "publicationTitle": "J",
            "volume": "",
            "issue": "",
            "pages": "",
        },
        "B": {
            "key": "B",
            "itemType": "journalArticle",
            "title": "Paper B",
            "creators": [],
            "date": "2024",
            "DOI": "",
            "publicationTitle": "J",
            "volume": "",
            "issue": "",
            "pages": "",
        },
    }
    result = build_document("Text [@A, @B].", item_data, "12345", str(output))
    doc = Document(result)
    xml = doc.paragraphs[0]._element.xml
    assert "ADDIN ZOTERO_ITEM" in xml


def test_build_bibliography_at_end(tmp_path):
    output = tmp_path / "biblio.docx"
    item_data = {
        "X": {
            "key": "X",
            "itemType": "journalArticle",
            "title": "Paper X",
            "creators": [],
            "date": "2024",
            "DOI": "",
            "publicationTitle": "J",
            "volume": "",
            "issue": "",
            "pages": "",
        },
    }
    result = build_document("Text [@X].", item_data, "12345", str(output))
    doc = Document(result)
    last_para = doc.paragraphs[-1]
    assert "ADDIN ZOTERO_BIBL" in last_para._element.xml


# -- insert_citations tests (in-place citation insertion) --

SAMPLE_ITEM_DATA = {
    "ABC": {
        "key": "ABC",
        "itemType": "journalArticle",
        "title": "Test Paper",
        "creators": [
            {"creatorType": "author", "firstName": "J", "lastName": "Doe"}
        ],
        "date": "2024",
        "DOI": "10.1/test",
        "publicationTitle": "Test Journal",
        "volume": "1",
        "issue": "1",
        "pages": "1-5",
    },
    "DEF": {
        "key": "DEF",
        "itemType": "journalArticle",
        "title": "Another Paper",
        "creators": [
            {"creatorType": "author", "firstName": "A", "lastName": "Smith"}
        ],
        "date": "2023",
        "DOI": "10.2/test",
        "publicationTitle": "Other Journal",
        "volume": "2",
        "issue": "4",
        "pages": "10-20",
    },
}


def _create_existing_doc(path, paragraphs, heading=None):
    """Helper: create a .docx with specific content for testing."""
    doc = Document()
    if heading:
        doc.add_heading(heading, level=1)
    for text, style in paragraphs:
        p = doc.add_paragraph(text)
        if style:
            p.style = doc.styles[style]
    doc.save(str(path))


def test_insert_citations_replaces_markers(tmp_path):
    doc_path = tmp_path / "existing.docx"
    _create_existing_doc(doc_path, [
        ("This is background text [@ABC].", None),
    ])

    result_path, count = insert_citations(
        str(doc_path), SAMPLE_ITEM_DATA, "12345"
    )
    assert count == 1

    doc = Document(result_path)
    xml = doc.paragraphs[0]._element.xml
    assert "ADDIN ZOTERO_ITEM" in xml
    # Original marker text should be gone
    plain_text = "".join(run.text for run in doc.paragraphs[0].runs)
    assert "[@ABC]" not in plain_text


def test_insert_citations_preserves_unmarked_paragraphs(tmp_path):
    doc_path = tmp_path / "mixed.docx"
    _create_existing_doc(doc_path, [
        ("This paragraph has no citations.", None),
        ("This one does [@ABC].", None),
        ("This one is also plain.", None),
    ])

    result_path, count = insert_citations(
        str(doc_path), SAMPLE_ITEM_DATA, "12345"
    )
    assert count == 1

    doc = Document(result_path)
    # First and third paragraphs should be untouched
    assert doc.paragraphs[0].text == "This paragraph has no citations."
    assert doc.paragraphs[2].text == "This one is also plain."
    # Second paragraph should have field codes
    assert "ADDIN ZOTERO_ITEM" in doc.paragraphs[1]._element.xml


def test_insert_citations_preserves_heading_style(tmp_path):
    doc_path = tmp_path / "styled.docx"
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Background text [@ABC].")
    doc.add_heading("Methods", level=2)
    doc.add_paragraph("We used standard methods.")
    doc.save(str(doc_path))

    result_path, _ = insert_citations(
        str(doc_path), SAMPLE_ITEM_DATA, "12345"
    )

    doc = Document(result_path)
    # Headings should be preserved
    assert doc.paragraphs[0].style.name == "Heading 1"
    assert doc.paragraphs[0].text == "Introduction"
    assert doc.paragraphs[2].style.name == "Heading 2"
    assert doc.paragraphs[2].text == "Methods"


def test_insert_citations_multiple_keys(tmp_path):
    doc_path = tmp_path / "multi.docx"
    _create_existing_doc(doc_path, [
        ("First ref [@ABC] and second ref [@DEF].", None),
    ])

    result_path, count = insert_citations(
        str(doc_path), SAMPLE_ITEM_DATA, "12345"
    )
    assert count == 2

    doc = Document(result_path)
    xml = doc.paragraphs[0]._element.xml
    assert xml.count("ADDIN ZOTERO_ITEM") == 2


def test_insert_citations_adds_bibliography(tmp_path):
    doc_path = tmp_path / "nobib.docx"
    _create_existing_doc(doc_path, [
        ("Some text [@ABC].", None),
    ])

    result_path, _ = insert_citations(
        str(doc_path), SAMPLE_ITEM_DATA, "12345"
    )

    doc = Document(result_path)
    last_para = doc.paragraphs[-1]
    assert "ADDIN ZOTERO_BIBL" in last_para._element.xml


def test_insert_citations_no_markers_returns_zero(tmp_path):
    doc_path = tmp_path / "plain.docx"
    _create_existing_doc(doc_path, [
        ("Just plain text with no markers.", None),
    ])

    result_path, count = insert_citations(
        str(doc_path), SAMPLE_ITEM_DATA, "12345"
    )
    assert count == 0

    # Document should be unchanged (no bibliography added)
    doc = Document(result_path)
    assert len(doc.paragraphs) == 1


def test_insert_citations_saves_to_different_path(tmp_path):
    doc_path = tmp_path / "original.docx"
    out_path = tmp_path / "output.docx"
    _create_existing_doc(doc_path, [
        ("Text [@ABC].", None),
    ])

    result_path, count = insert_citations(
        str(doc_path), SAMPLE_ITEM_DATA, "12345", str(out_path)
    )
    assert count == 1
    assert Path(result_path) == out_path.resolve()
    assert out_path.exists()


def test_insert_citations_table_cells(tmp_path):
    """Citation markers inside table cells should also be replaced."""
    doc_path = tmp_path / "table.docx"
    doc = Document()
    doc.add_paragraph("Normal paragraph.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cell with ref [@ABC]."
    table.cell(0, 1).text = "Plain cell."
    doc.save(str(doc_path))

    result_path, count = insert_citations(
        str(doc_path), SAMPLE_ITEM_DATA, "12345"
    )
    assert count == 1

    doc = Document(result_path)
    cell_xml = doc.tables[0].cell(0, 0).paragraphs[0]._element.xml
    assert "ADDIN ZOTERO_ITEM" in cell_xml
